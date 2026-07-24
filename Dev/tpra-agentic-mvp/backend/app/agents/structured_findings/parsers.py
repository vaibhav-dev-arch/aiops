"""Input parsers for XLSX, CSV, JSON, DOCX."""

from __future__ import annotations

import csv
import io
import json
from typing import Any

from docx import Document
from openpyxl import load_workbook


def parse_bytes(data: bytes, filename: str) -> list[dict[str, Any]]:
    name = filename.lower()
    if name.endswith(".csv"):
        return _parse_csv(data)
    if name.endswith(".json"):
        return _parse_json(data)
    if name.endswith((".xlsx", ".xls")):
        return _parse_xlsx(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    # fallback: try CSV then JSON
    try:
        return _parse_csv(data)
    except Exception:
        return _parse_json(data)


def _parse_csv(data: bytes) -> list[dict[str, Any]]:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = []
    for row in reader:
        cleaned = {str(k).strip(): ("" if v is None else str(v).strip()) for k, v in row.items() if k}
        if any(cleaned.values()):
            rows.append(cleaned)
    return rows


def _parse_json(data: bytes) -> list[dict[str, Any]]:
    parsed = json.loads(data.decode("utf-8"))
    if isinstance(parsed, dict):
        if "findings" in parsed and isinstance(parsed["findings"], list):
            return [dict(x) for x in parsed["findings"]]
        return [parsed]
    if isinstance(parsed, list):
        return [dict(x) for x in parsed]
    raise ValueError("JSON must be an object or array of findings")


def _parse_xlsx(data: bytes) -> list[dict[str, Any]]:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)
    try:
        headers = [str(h).strip() if h is not None else f"col_{i}" for i, h in enumerate(next(rows_iter))]
    except StopIteration:
        return []
    out: list[dict[str, Any]] = []
    for row in rows_iter:
        item = {
            headers[i]: ("" if (i >= len(row) or row[i] is None) else str(row[i]).strip())
            for i in range(len(headers))
        }
        if any(item.values()):
            out.append(item)
    return out


def _parse_docx(data: bytes) -> list[dict[str, Any]]:
    doc = Document(io.BytesIO(data))
    # Prefer first table if present
    if doc.tables:
        table = doc.tables[0]
        cells = table.rows[0].cells
        headers = [c.text.strip() or f"col_{i}" for i, c in enumerate(cells)]
        out: list[dict[str, Any]] = []
        for row in table.rows[1:]:
            values = [c.text.strip() for c in row.cells]
            item = {headers[i]: (values[i] if i < len(values) else "") for i in range(len(headers))}
            if any(item.values()):
                out.append(item)
        if out:
            return out
    # Fallback: each non-empty paragraph becomes a finding title
    findings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            findings.append({"title": text, "description": "", "severity": "medium", "category": "general"})
    return findings
