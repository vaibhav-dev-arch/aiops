"""DOCX TPRA report builder."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt


def build_report_docx(
    *,
    title: str,
    narrative: str,
    sections: dict[str, list[dict[str, Any]]],
    missing_inputs: list[str] | None = None,
) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph("Status: Draft — for human review")
    doc.add_heading("AI-Generated Narrative", level=1)
    for para in narrative.split("\n\n"):
        text = para.strip()
        if text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_heading("Findings by Section", level=1)
    for section_name, findings in sections.items():
        doc.add_heading(section_name, level=2)
        for f in findings:
            doc.add_paragraph(
                f"[{f.get('severity', '').upper()}] {f.get('title', '')} — {f.get('description', '')}",
                style="List Bullet",
            )

    if missing_inputs:
        doc.add_heading("Missing Inputs", level=1)
        for item in missing_inputs:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph(
        "Reviewer to validate narrative accuracy, confirm residual risk ratings, and approve for finalization."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
